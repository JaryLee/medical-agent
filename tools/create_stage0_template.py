from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parents[1] / "backend/src/main/resources/templates/anonymous-research-protocol.docx"

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
]:
    style = doc.styles[name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.text = "医疗研究 Agent · 阶段 0 匿名原型"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.runs[0].font.name = "Microsoft YaHei"
header.runs[0].font.size = Pt(9)
header.runs[0].font.color.rgb = RGBColor(100, 100, 100)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(5)
run = title.add_run("${project.title}")
run.bold = True
run.font.name = "Microsoft YaHei"
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(31, 77, 120)

subtitle = doc.add_paragraph("匿名技术原型 · 非正式研究方案")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.name = "Microsoft YaHei"
subtitle.runs[0].font.size = Pt(10)
subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

doc.add_heading("研究概览", level=1)
table = doc.add_table(rows=3, cols=2)
table.autofit = False
table.columns[0].width = Inches(1.875)
table.columns[1].width = Inches(4.625)
rows = [
    ("研究问题", "${research.question}"),
    ("建议设计", "${research.studyDesign}"),
    ("研究人群", "${research.population}"),
]
for row, values in zip(table.rows, rows):
    row.cells[0].width = Inches(1.875)
    row.cells[1].width = Inches(4.625)
    for cell, text in zip(row.cells, values):
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(2)
            for cell_run in paragraph.runs:
                cell_run.font.name = "Microsoft YaHei"
                cell_run.font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].bold = True

tbl_pr = table._tbl.tblPr
tbl_w = tbl_pr.find(qn("w:tblW"))
tbl_w.set(qn("w:type"), "dxa")
tbl_w.set(qn("w:w"), "9360")
tbl_ind = OxmlElement("w:tblInd")
tbl_ind.set(qn("w:w"), "120")
tbl_ind.set(qn("w:type"), "dxa")
tbl_pr.append(tbl_ind)

doc.add_heading("研究背景", level=1)
doc.add_paragraph("${research.background}")
doc.add_heading("主要结局", level=1)
doc.add_paragraph("${research.outcomes}")
doc.add_heading("参考文献", level=1)
doc.add_paragraph("${research.references}")
note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(14)
note_run = note.add_run("重要说明：本文件由确定性 Mock 数据生成，仅用于技术验证；不能替代医学、统计学、伦理或科研管理专家审核。")
note_run.bold = True
note_run.font.name = "Microsoft YaHei"
note_run.font.color.rgb = RGBColor(122, 90, 0)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
