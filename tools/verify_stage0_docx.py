from pathlib import Path
from docx import Document

path = Path(__file__).resolve().parents[1] / "artifacts/阶段0-匿名研究方案.docx"
document = Document(path)
text = "\n".join(p.text for p in document.paragraphs)
text += "\n" + "\n".join(
    cell.text
    for table in document.tables
    for row in table.rows
    for cell in row.cells
)

checks = {
    "no_unresolved_placeholders": "${" not in text,
    "contains_selected_design": "回顾性队列研究" in text,
    "contains_verified_mock_pmid": "36331190" in text,
    "contains_mock_disclaimer": "确定性 Mock 数据" in text,
    "has_research_overview_table": len(document.tables) == 1,
}

failed = [name for name, passed in checks.items() if not passed]
print({"paragraphs": len(document.paragraphs), "tables": len(document.tables), **checks})
if failed:
    raise SystemExit("DOCX structural checks failed: " + ", ".join(failed))
