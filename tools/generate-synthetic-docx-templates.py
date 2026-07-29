from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\develop\AIWorkspace\MEDICAL_AGENT")
OUTPUT = ROOT / "backend" / "src" / "test" / "resources" / "templates"
OUTPUT.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
NAVY = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = RGBColor(95, 99, 104)
BLACK = RGBColor(0, 0, 0)


def set_run_font(run, size=10.5, bold=False, color=BLACK, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def configure_section_geometry(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def configure_document(document, preset):
    configure_section_geometry(document.sections[0])

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5 if preset == "narrative" else 10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8 if preset == "narrative" else 6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    document.core_properties.creator = "Medical Research Agent"
    document.core_properties.subject = "SYNTHETIC_ANONYMOUS test template"
    document.core_properties.keywords = (
        "SYNTHETIC_ANONYMOUS; NON_HOSPITAL_MATERIAL; TEST_ONLY"
    )


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, vertical_margin=80):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    grid_cols = list(grid)
    while len(grid_cols) < len(widths):
        node = OxmlElement("w:gridCol")
        grid.append(node)
        grid_cols.append(node)
    for node, width in zip(grid_cols, widths):
        node.set(qn("w:w"), str(width))

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, top=vertical_margin, bottom=vertical_margin)


def set_cell_text(cell, text, label=False, size=10):
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(
        run,
        size=size,
        bold=label,
        color=RGBColor.from_string(NAVY) if label else BLACK,
    )


def add_body(document, text, after=8, justified=True):
    paragraph = document.add_paragraph()
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY if justified else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_header_footer(document, label):
    header = document.sections[0].header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run(f"{label} · 非真实医院材料")
    set_run_font(run, size=8.5, color=MUTED)

    footer = document.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_after = Pt(0)
    run = footer_p.add_run("SYNTHETIC_ANONYMOUS · 仅用于模板兼容与回归测试")
    set_run_font(run, size=8, color=MUTED)


def build_template_a():
    document = Document()
    configure_document(document, "narrative")
    add_header_footer(document, "合成测试模板 A")
    document.core_properties.title = "合成测试模板 A：叙事型观察性研究方案"

    cover_label = document.add_paragraph()
    cover_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cover_label.paragraph_format.space_before = Pt(24)
    cover_label.paragraph_format.space_after = Pt(8)
    set_run_font(
        cover_label.add_run("合成测试模板 A · 非真实医院材料"),
        size=8.5,
        color=MUTED,
    )

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(36)
    title.paragraph_format.space_after = Pt(6)
    set_run_font(title.add_run("${project.title}"), size=24, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(
        subtitle.add_run("观察性研究方案 · 合成测试版"),
        size=13,
        color=MUTED,
    )

    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2160, 7200])
    rows = [
        ("课题负责人", "${project.principalInvestigator}"),
        ("所属科室", "${project.department}"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_cell_text(row.cells[0], label, label=True)
        set_cell_text(row.cells[1], value)

    document.add_heading("1. 研究背景与问题", level=1)
    add_body(document, "${research.background}")
    document.add_heading("1.1 核心研究问题", level=2)
    add_body(document, "${research.question}")
    document.add_heading("1.2 研究目标", level=2)
    add_body(document, "${research.objectives}")

    document.add_heading("2. 研究设计", level=1)
    add_body(document, "${research.studyDesign}")
    document.add_heading("2.1 研究对象", level=2)
    add_body(document, "${research.population}")
    document.add_heading("2.2 纳入标准", level=2)
    add_body(document, "${research.inclusionCriteria}")
    document.add_heading("2.3 排除标准", level=2)
    add_body(document, "${research.exclusionCriteria}")

    analysis_label = document.add_paragraph()
    analysis_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    analysis_label.paragraph_format.keep_with_next = True
    analysis_label.paragraph_format.space_before = Pt(24)
    analysis_label.paragraph_format.space_after = Pt(16)
    set_run_font(
        analysis_label.add_run("分析与引用附页 · SYNTHETIC_ANONYMOUS"),
        size=8.5,
        color=MUTED,
    )

    document.add_heading("3. 终点、变量与统计", level=1)
    document.add_heading("3.1 研究终点", level=2)
    add_body(document, "${research.outcomes}")
    document.add_heading("3.2 变量定义", level=2)
    add_body(document, "${research.variables}")
    document.add_heading("3.3 统计分析计划", level=2)
    add_body(document, "${research.statisticalPlan}")

    document.add_heading("4. 伦理与数据安全", level=1)
    add_body(document, "${research.ethicalConsiderations}")
    document.add_heading("5. 参考文献", level=1)
    add_body(document, "${research.references}", justified=False)

    warning = document.add_paragraph()
    warning.paragraph_format.space_before = Pt(12)
    warning.paragraph_format.space_after = Pt(0)
    set_run_font(
        warning.add_run(
            "重要说明：本模板为系统生成的合成测试材料，不代表任何真实医院格式。"
        ),
        size=9,
        bold=True,
        color=RGBColor(128, 85, 0),
    )
    output = OUTPUT / "synthetic-hospital-a-protocol.docx"
    document.save(output)
    return output


def add_form_row(table, label, placeholder, fill=LIGHT_GRAY):
    row = table.add_row()
    set_table_geometry(table, [2160, 7200], vertical_margin=35)
    set_cell_shading(row.cells[0], fill)
    set_cell_text(row.cells[0], label, label=True)
    set_cell_text(row.cells[1], placeholder)


def build_template_b():
    document = Document()
    configure_document(document, "compact")
    add_header_footer(document, "合成测试模板 B")
    document.core_properties.title = "合成测试模板 B：表格式科研方案送审稿"

    logo = document.add_paragraph()
    # LibreOffice needs explicit first-body spacing to keep page furniture
    # outside the running header on page one.
    logo.paragraph_format.space_before = Pt(72)
    logo.paragraph_format.space_after = Pt(4)
    set_run_font(logo.add_run("${project.logo}"), size=10)

    kicker = document.add_paragraph()
    # Explicit paragraph spacing survives the Apache POI template round-trip
    # and keeps LibreOffice from laying the first body paragraph over the
    # first-page running header.
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(2)
    set_run_font(
        kicker.add_run("SYNTHETIC RESEARCH SUBMISSION"),
        size=9,
        bold=True,
        color=RGBColor.from_string(BLUE),
    )
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("医学科研方案送审稿"), size=22, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(
        subtitle.add_run("合成测试模板 B · 非真实医院材料"),
        size=11,
        color=MUTED,
    )

    metadata = document.add_table(rows=0, cols=2)
    metadata.style = "Table Grid"
    add_form_row(metadata, "课题名称", "${project.title}", LIGHT_BLUE)
    add_form_row(metadata, "课题负责人", "${project.principalInvestigator}", LIGHT_BLUE)
    add_form_row(metadata, "所属科室", "${project.department}", LIGHT_BLUE)
    add_form_row(metadata, "研究设计", "${research.studyDesign}", LIGHT_BLUE)

    document.add_heading("一、研究摘要", level=1)
    summary = document.add_table(rows=0, cols=2)
    summary.style = "Table Grid"
    add_form_row(summary, "研究背景", "${research.background}")
    add_form_row(summary, "研究问题", "${research.question}")
    add_form_row(summary, "研究目标", "${list.research.objectives}")
    add_form_row(summary, "研究对象", "${research.population}")

    document.add_heading("二、研究对象与变量", level=1)
    criteria = document.add_table(rows=0, cols=2)
    criteria.style = "Table Grid"
    add_form_row(criteria, "纳入标准", "${research.inclusionCriteria}")
    add_form_row(criteria, "排除标准", "${research.exclusionCriteria}")
    add_form_row(criteria, "主要与次要终点", "${research.outcomes}")
    add_form_row(criteria, "变量与混杂因素", "${research.variables}")

    continuation_label = document.add_paragraph()
    continuation_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    continuation_label.paragraph_format.page_break_before = True
    continuation_label.paragraph_format.space_before = Pt(24)
    continuation_label.paragraph_format.space_after = Pt(16)
    set_run_font(
        continuation_label.add_run(
            "分析与引用附页 · SYNTHETIC_ANONYMOUS"
        ),
        size=8.5,
        color=MUTED,
    )

    document.add_heading("三、分析与合规", level=1)
    compliance = document.add_table(rows=0, cols=2)
    compliance.style = "Table Grid"
    add_form_row(compliance, "统计分析计划", "${research.statisticalPlan}")
    add_form_row(
        compliance,
        "伦理与数据安全",
        "${research.ethicalConsiderations}",
    )

    references_heading = document.add_heading("四、参考文献", level=1)
    references_heading.paragraph_format.space_before = Pt(12)
    references = document.add_table(rows=2, cols=2)
    references.style = "Table Grid"
    set_table_geometry(references, [1600, 7760], vertical_margin=90)
    set_cell_shading(references.rows[0].cells[0], LIGHT_BLUE)
    set_cell_shading(references.rows[0].cells[1], LIGHT_BLUE)
    set_cell_text(references.rows[0].cells[0], "序号", label=True)
    set_cell_text(references.rows[0].cells[1], "经核验的参考文献", label=True)
    set_cell_text(
        references.rows[1].cells[0],
        "${repeat.ref.index}",
        size=8,
    )
    set_cell_text(
        references.rows[1].cells[1],
        "${repeat.ref.text}",
    )

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    set_run_font(
        note.add_run(
            "模板状态：SYNTHETIC_ANONYMOUS / TEST_ONLY。真实医院模板到位后必须重新适配与验收。"
        ),
        size=9,
        bold=True,
        color=RGBColor(128, 85, 0),
    )
    output = OUTPUT / "synthetic-hospital-b-protocol.docx"
    document.save(output)
    return output


if __name__ == "__main__":
    for path in (build_template_a(), build_template_b()):
        print(f"{path} {path.stat().st_size}")
